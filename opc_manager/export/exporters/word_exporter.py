import io

from docx import Document


class WordExporter:
    def export(self, data, template=None, **opts) -> bytes:
        doc = Document()
        title = data.metadata.get("title", "Document")
        doc.add_heading(title, level=0)

        for line in data.content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif "| " in stripped:
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                if any(c.replace("-", "") == "" for c in cells):
                    continue
                table = doc.add_table(rows=1, cols=len(cells))
                for i, cell in enumerate(cells):
                    table.rows[0].cells[i].text = cell
            else:
                doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
