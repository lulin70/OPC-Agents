"""
OPC-Agents 文档处理技能

功能：
- PDF 文档解析（文本提取、元数据）
- Word 文档处理（.docx 读写）
- Excel 表格处理（读取、计算、基础分析）
"""

import os
from typing import Dict, List, Optional, Any
from pathlib import Path


class DocumentProcessorSkill:
    """文档处理技能"""
    
    # 技能元数据
    METADATA = {
        'name': 'document_processor',
        'version': '1.0.0',
        'description': '多功能文档处理技能，支持 PDF/Word/Excel',
        'author': 'OPC-Agents Team',
        'category': 'document_processing',
        'tags': ['PDF', 'Word', 'Excel', '文档', '办公'],
        'permissions': ['read_file', 'write_file'],
    }
    
    def __init__(self, config: Optional[Dict] = None):
        """
        初始化文档处理技能
        
        Args:
            config: 配置字典
        """
        self.config = config or {}
        self.output_dir = self.config.get('output_dir', './output')
        
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
    
    def execute(self, 
                operation: str,
                file_path: str,
                **kwargs) -> Dict:
        """
        执行文档操作
        
        Args:
            operation: 操作类型 ('read_pdf', 'read_word', 'read_excel', 'write_word', 'write_excel')
            file_path: 文件路径
            **kwargs: 其他参数
            
        Returns:
            Dict: 操作结果
        """
        try:
            if operation == 'read_pdf':
                return self._read_pdf(file_path, **kwargs)
            elif operation == 'read_word':
                return self._read_word(file_path, **kwargs)
            elif operation == 'read_excel':
                return self._read_excel(file_path, **kwargs)
            elif operation == 'write_word':
                return self._write_word(file_path, **kwargs)
            elif operation == 'write_excel':
                return self._write_excel(file_path, **kwargs)
            else:
                return {
                    'success': False,
                    'error': f'不支持的操作：{operation}',
                    'supported_operations': ['read_pdf', 'read_word', 'read_excel', 'write_word', 'write_excel']
                }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
            }
    
    def _read_pdf(self, 
                  file_path: str, 
                  extract_images: bool = False,
                  **kwargs) -> Dict:
        """
        读取 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            extract_images: 是否提取图片
            **kwargs: 其他参数
            
        Returns:
            Dict: {
                'success': bool,
                'text': str,
                'pages': int,
                'metadata': Dict,
                'images': List (if extract_images)
            }
        """
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                # 提取文本
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                # 提取元数据
                metadata = {
                    'pages': len(pdf.pages),
                    'author': pdf.metadata.get('Author', '') if pdf.metadata else '',
                    'title': pdf.metadata.get('Title', '') if pdf.metadata else '',
                    'subject': pdf.metadata.get('Subject', '') if pdf.metadata else '',
                    'creator': pdf.metadata.get('Creator', '') if pdf.metadata else '',
                }
                
                # 提取图片（可选）
                images = []
                if extract_images:
                    for i, page in enumerate(pdf.pages):
                        page_images = page.images
                        images.append({
                            'page': i + 1,
                            'count': len(page_images)
                        })
                
                return {
                    'success': True,
                    'text': text,
                    'pages': metadata['pages'],
                    'metadata': metadata,
                    'images': images if extract_images else None,
                    'file_path': file_path,
                }
                
        except ImportError:
            return {
                'success': False,
                'error': '缺少依赖：pdfplumber，请运行 pip3 install pdfplumber'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'读取 PDF 失败：{str(e)}'
            }
    
    def _read_word(self, 
                   file_path: str,
                   extract_tables: bool = False,
                   **kwargs) -> Dict:
        """
        读取 Word 文档
        
        Args:
            file_path: Word 文件路径 (.docx)
            extract_tables: 是否提取表格
            **kwargs: 其他参数
            
        Returns:
            Dict: {
                'success': bool,
                'content': str,
                'paragraphs': List[str],
                'tables': List (if extract_tables),
                'metadata': Dict
            }
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 合并内容
            content = "\n\n".join(paragraphs)
            
            # 提取表格（可选）
            tables = []
            if extract_tables:
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append({
                        'index': i,
                        'rows': len(table.rows),
                        'cols': len(table.columns),
                        'data': table_data
                    })
            
            # 元数据
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
            }
            
            return {
                'success': True,
                'content': content,
                'paragraphs': paragraphs,
                'tables': tables if extract_tables else None,
                'metadata': metadata,
                'file_path': file_path,
            }
            
        except ImportError:
            return {
                'success': False,
                'error': '缺少依赖：python-docx，请运行 pip3 install python-docx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'读取 Word 文档失败：{str(e)}'
            }
    
    def _read_excel(self, 
                    file_path: str,
                    sheet_name: Optional[str] = None,
                    include_formulas: bool = False,
                    **kwargs) -> Dict:
        """
        读取 Excel 文件
        
        Args:
            file_path: Excel 文件路径 (.xlsx/.xls)
            sheet_name: 工作表名称（默认第一个）
            include_formulas: 是否包含公式
            **kwargs: 其他参数
            
        Returns:
            Dict: {
                'success': bool,
                'data': List[Dict],
                'sheets': List[str],
                'shape': Tuple,
                'columns': List[str]
            }
        """
        try:
            import pandas as pd
            from openpyxl import load_workbook
            
            # 读取所有工作表名称
            wb = load_workbook(file_path, data_only=not include_formulas)
            sheets = wb.sheetnames
            wb.close()
            
            # 读取指定工作表
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                df = pd.read_excel(file_path, sheet_name=0)
                sheet_name = sheets[0]
            
            # 转换为字典列表
            data = df.to_dict('records')
            
            # 处理 NaN 值
            for row in data:
                for key in row:
                    if pd.isna(row[key]):
                        row[key] = None
            
            return {
                'success': True,
                'data': data,
                'sheets': sheets,
                'current_sheet': sheet_name,
                'shape': list(df.shape),  # [行数，列数]
                'columns': list(df.columns),
                'file_path': file_path,
            }
            
        except ImportError:
            return {
                'success': False,
                'error': '缺少依赖：pandas 或 openpyxl，请运行 pip3 install pandas openpyxl'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'读取 Excel 文件失败：{str(e)}'
            }
    
    def _write_word(self,
                    file_path: str,
                    content: str,
                    title: Optional[str] = None,
                    paragraphs: Optional[List[str]] = None,
                    **kwargs) -> Dict:
        """
        创建/写入 Word 文档
        
        Args:
            file_path: 输出文件路径
            content: 文档内容
            title: 文档标题
            paragraphs: 段落列表（可选）
            **kwargs: 其他参数
            
        Returns:
            Dict: {
                'success': bool,
                'file_path': str
            }
        """
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # 添加标题
            if title:
                doc.add_heading(title, level=1)
            
            # 添加内容
            if paragraphs:
                for para in paragraphs:
                    doc.add_paragraph(para)
            elif content:
                # 按空行分割内容
                para_list = content.split('\n\n')
                for para in para_list:
                    if para.strip():
                        doc.add_paragraph(para)
            
            # 保存文件
            doc.save(file_path)
            
            return {
                'success': True,
                'file_path': file_path,
                'message': f'Word 文档已创建：{file_path}'
            }
            
        except ImportError:
            return {
                'success': False,
                'error': '缺少依赖：python-docx，请运行 pip3 install python-docx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'创建 Word 文档失败：{str(e)}'
            }
    
    def _write_excel(self,
                     file_path: str,
                     data: List[Dict],
                     sheet_name: str = 'Sheet1',
                     **kwargs) -> Dict:
        """
        创建/写入 Excel 文件
        
        Args:
            file_path: 输出文件路径
            data: 数据列表（字典列表）
            sheet_name: 工作表名称
            **kwargs: 其他参数
            
        Returns:
            Dict: {
                'success': bool,
                'file_path': str
            }
        """
        try:
            import pandas as pd
            
            # 转换为 DataFrame
            df = pd.DataFrame(data)
            
            # 写入 Excel
            df.to_excel(file_path, sheet_name=sheet_name, index=False)
            
            return {
                'success': True,
                'file_path': file_path,
                'message': f'Excel 文件已创建：{file_path}',
                'rows': len(data),
                'columns': len(data[0]) if data else 0
            }
            
        except ImportError:
            return {
                'success': False,
                'error': '缺少依赖：pandas 和 openpyxl，请运行 pip3 install pandas openpyxl'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'创建 Excel 文件失败：{str(e)}'
            }
    
    def get_schema(self) -> Dict:
        """返回输入输出 schema"""
        return {
            'input': {
                'operation': {
                    'type': 'string',
                    'required': True,
                    'description': '操作类型',
                    'enum': ['read_pdf', 'read_word', 'read_excel', 'write_word', 'write_excel']
                },
                'file_path': {
                    'type': 'string',
                    'required': True,
                    'description': '文件路径'
                },
                'extract_images': {
                    'type': 'boolean',
                    'required': False,
                    'description': '是否提取 PDF 图片'
                },
                'extract_tables': {
                    'type': 'boolean',
                    'required': False,
                    'description': '是否提取 Word 表格'
                },
                'sheet_name': {
                    'type': 'string',
                    'required': False,
                    'description': 'Excel 工作表名称'
                },
                'content': {
                    'type': 'string',
                    'required': False,
                    'description': '写入 Word 的内容'
                },
                'data': {
                    'type': 'array',
                    'required': False,
                    'description': '写入 Excel 的数据'
                },
            },
            'output': {
                'success': {'type': 'boolean'},
                'text': {'type': 'string'},  # PDF 文本
                'content': {'type': 'string'},  # Word 内容
                'data': {'type': 'array'},  # Excel 数据
                'paragraphs': {'type': 'array'},
                'tables': {'type': 'array'},
                'metadata': {'type': 'object'},
                'file_path': {'type': 'string'},
                'error': {'type': 'string'},
            }
        }


# 便捷函数
def read_pdf(file_path: str, **kwargs) -> Dict:
    """读取 PDF 文件"""
    skill = DocumentProcessorSkill()
    return skill.execute('read_pdf', file_path, **kwargs)


def read_word(file_path: str, **kwargs) -> Dict:
    """读取 Word 文档"""
    skill = DocumentProcessorSkill()
    return skill.execute('read_word', file_path, **kwargs)


def read_excel(file_path: str, **kwargs) -> Dict:
    """读取 Excel 文件"""
    skill = DocumentProcessorSkill()
    return skill.execute('read_excel', file_path, **kwargs)


def write_word(file_path: str, content: str, **kwargs) -> Dict:
    """创建 Word 文档"""
    skill = DocumentProcessorSkill()
    return skill.execute('write_word', file_path, content=content, **kwargs)


def write_excel(file_path: str, data: List[Dict], **kwargs) -> Dict:
    """创建 Excel 文件"""
    skill = DocumentProcessorSkill()
    return skill.execute('write_excel', file_path, data=data, **kwargs)


# 测试
if __name__ == '__main__':
    print("=" * 60)
    print("文档处理技能测试")
    print("=" * 60)
    
    skill = DocumentProcessorSkill()
    
    # 测试 schema
    print("\n1. 测试获取 schema")
    schema = skill.get_schema()
    print(f"✅ Schema 获取成功，包含 {len(schema['input'])} 个输入参数")
    
    # 测试不支持的操作
    print("\n2. 测试不支持的操作")
    result = skill.execute('invalid_operation', 'test.txt')
    print(f"❌ 操作结果：{result['error']}")
    
    # 测试创建 Word 文档
    print("\n3. 测试创建 Word 文档")
    test_content = """这是第一段内容。

这是第二段内容。

这是第三段内容。"""
    
    output_word = os.path.join(skill.output_dir, 'test_output.docx')
    result = skill.execute(
        'write_word',
        output_word,
        content=test_content,
        title='测试文档'
    )
    
    if result['success']:
        print(f"✅ Word 文档创建成功：{result['file_path']}")
    else:
        print(f"❌ Word 文档创建失败：{result['error']}")
    
    # 测试创建 Excel 文档
    print("\n4. 测试创建 Excel 文档")
    test_data = [
        {'姓名': '张三', '年龄': 25, '城市': '北京'},
        {'姓名': '李四', '年龄': 30, '城市': '上海'},
        {'姓名': '王五', '年龄': 28, '城市': '广州'},
    ]
    
    output_excel = os.path.join(skill.output_dir, 'test_output.xlsx')
    result = skill.execute(
        'write_excel',
        output_excel,
        data=test_data
    )
    
    if result['success']:
        print(f"✅ Excel 文件创建成功：{result['file_path']}")
        print(f"   行数：{result.get('rows', 0)}, 列数：{result.get('columns', 0)}")
    else:
        print(f"❌ Excel 文件创建失败：{result['error']}")
    
    # 测试读取 Excel
    print("\n5. 测试读取 Excel 文件")
    if os.path.exists(output_excel):
        result = skill.execute('read_excel', output_excel)
        if result['success']:
            print(f"✅ Excel 文件读取成功")
            print(f"   工作表：{result['sheets']}")
            print(f"   形状：{result['shape']}")
            print(f"   列：{result['columns']}")
            print(f"   前 2 行数据:")
            for i, row in enumerate(result['data'][:2], 1):
                print(f"     {i}. {row}")
        else:
            print(f"❌ Excel 文件读取失败：{result['error']}")
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)
